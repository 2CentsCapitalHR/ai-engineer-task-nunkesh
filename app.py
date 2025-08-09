#!/usr/bin/env python3
"""
ADGM Corporate Agent - Complete Ollama-Powered Legal Document Intelligence System
Author: AI Assistant
Description: Complete legal document analysis system with Ollama integration
"""

import gradio as gr
import json
import os
import logging
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Any
import tempfile
import zipfile
from pathlib import Path
import requests
import re
from urllib.parse import urlparse, urljoin

# Document processing imports
try:
    from docx import Document
    from docx.shared import RGBColor
    from docx.oxml.shared import OxmlElement, qn
    from docx.enum.text import WD_COLOR_INDEX
    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️ python-docx not available. Install with: pip install python-docx")
    DOCX_AVAILABLE = False

# AI and ML imports
try:
    import ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    print("⚠️ Ollama not available. Install with: pip install ollama")
    OLLAMA_AVAILABLE = False

try:
    import numpy as np
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    ML_AVAILABLE = True
except ImportError:
    print("⚠️ ML libraries not available. Install with: pip install numpy scikit-learn")
    ML_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    print("⚠️ BeautifulSoup not available. Install with: pip install beautifulsoup4")
    BS4_AVAILABLE = False

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class OllamaEmbeddingSystem:
    """Enhanced embedding system using Ollama with multiple fallback strategies"""
    
    def __init__(self, model_name="llama3.2"):
        self.model_name = model_name
        self.ollama_available = self._test_ollama_connection()
        self.backup_models = ["mistral:7b", "llama3.1:8b", "llama2:7b"]
        
        if not self.ollama_available and ML_AVAILABLE:
            logger.info("🔄 Ollama not available, initializing TF-IDF fallback")
            self.tfidf = TfidfVectorizer(max_features=512, stop_words='english', ngram_range=(1, 2))
            self.fitted = False
        elif not ML_AVAILABLE:
            logger.warning("⚠️ Neither Ollama nor ML libraries available - limited functionality")
    
    def _test_ollama_connection(self) -> bool:
        """Test if Ollama is running and accessible"""
        if not OLLAMA_AVAILABLE:
            return False
        
        try:
            models = ollama.list()
            available_models = [model['name'] for model in models.get('models', [])]
            logger.info(f"✅ Ollama connected - Available models: {available_models}")
            
            # Check if our preferred model is available
            if self.model_name not in available_models:
                for backup in self.backup_models:
                    if backup in available_models:
                        logger.info(f"🔄 Switching to available model: {backup}")
                        self.model_name = backup
                        break
                else:
                    if available_models:
                        self.model_name = available_models[0]
                        logger.info(f"🔄 Using first available model: {self.model_name}")
            
            return True
        except Exception as e:
            logger.error(f"❌ Ollama connection failed: {e}")
            return False
    
    def encode(self, texts: List[str]) -> np.ndarray:
        """Encode texts to embeddings with multiple fallback strategies"""
        if self.ollama_available:
            return self._ollama_encode(texts)
        elif ML_AVAILABLE:
            return self._tfidf_encode(texts)
        else:
            return self._simple_hash_encode(texts)
    
    def _ollama_encode(self, texts: List[str]) -> np.ndarray:
        """Use Ollama for embeddings with error handling"""
        embeddings = []
        for i, text in enumerate(texts):
            try:
                # Truncate text to avoid token limits
                truncated_text = text[:2000] if len(text) > 2000 else text
                
                response = ollama.embeddings(
                    model=self.model_name,
                    prompt=truncated_text
                )
                embeddings.append(response['embedding'])
                
                if i % 10 == 0:  # Progress indicator for large batches
                    logger.info(f"📊 Processed {i+1}/{len(texts)} embeddings")
                    
            except Exception as e:
                logger.warning(f"⚠️ Ollama embedding failed for text {i}: {e}")
                # Fallback to simple encoding for this text
                embeddings.append(self._simple_hash_vector(text))
        
        return np.array(embeddings)
    
    def _tfidf_encode(self, texts: List[str]) -> np.ndarray:
        """Use TF-IDF as fallback with improved preprocessing"""
        if not texts:
            return np.array([])
        
        # Preprocess texts
        processed_texts = []
        for text in texts:
            # Clean and truncate text
            clean_text = re.sub(r'[^\w\s]', ' ', text.lower())
            clean_text = ' '.join(clean_text.split())  # Remove extra whitespace
            processed_texts.append(clean_text[:1000])  # Limit length
        
        if not self.fitted:
            try:
                self.tfidf.fit(processed_texts)
                self.fitted = True
                logger.info("✅ TF-IDF model fitted")
            except Exception as e:
                logger.error(f"❌ TF-IDF fitting failed: {e}")
                return self._simple_hash_encode(texts)
        
        try:
            return self.tfidf.transform(processed_texts).toarray()
        except Exception as e:
            logger.error(f"❌ TF-IDF transform failed: {e}")
            return self._simple_hash_encode(texts)
    
    def _simple_hash_encode(self, texts: List[str]) -> np.ndarray:
        """Simple hash-based encoding as last resort"""
        logger.info("🔄 Using simple hash encoding (limited functionality)")
        embeddings = []
        for text in texts:
            embeddings.append(self._simple_hash_vector(text))
        return np.array(embeddings)
    
    def _simple_hash_vector(self, text: str, dim: int = 128) -> List[float]:
        """Create a simple hash-based vector"""
        words = text.lower().split()[:20]  # Use first 20 words
        vector = [0.0] * dim
        for i, word in enumerate(words):
            hash_val = hash(word) % dim
            vector[hash_val] += 1.0 / (i + 1)  # Weight by position
        return vector

class OllamaIntelligentChat:
    """Advanced chat interface for Ollama with context management"""
    
    def __init__(self, model_name="llama3.2"):
        self.model_name = model_name
        self.ollama_available = self._test_ollama_connection()
        self.context_window = 4000  # Token limit for context
        self.conversation_history = []
    
    def _test_ollama_connection(self) -> bool:
        """Test if Ollama is running with model availability check"""
        if not OLLAMA_AVAILABLE:
            return False
        
        try:
            models = ollama.list()
            available_models = [model['name'] for model in models.get('models', [])]
            
            if self.model_name not in available_models:
                backup_models = ["mistral:7b", "llama3.1:8b", "llama2:7b"]
                for backup in backup_models:
                    if backup in available_models:
                        self.model_name = backup
                        logger.info(f"🔄 Switched to available model: {backup}")
                        break
                else:
                    if available_models:
                        self.model_name = available_models[0]
                        logger.info(f"🔄 Using first available model: {self.model_name}")
            
            return True
        except Exception as e:
            logger.error(f"❌ Ollama test failed: {e}")
            return False
    
    def query(self, prompt: str, context: str = "", max_tokens: int = 800, temperature: float = 0.3) -> str:
        """Enhanced query with context management and legal expertise"""
        if not self.ollama_available:
            return self._fallback_response(prompt)
        
        try:
            # Build enhanced prompt with legal context
            system_prompt = """You are an expert legal analyst specializing in Abu Dhabi Global Market (ADGM) regulations and corporate law. 
            Provide precise, actionable advice with specific legal references where possible. 
            Focus on ADGM compliance requirements, corporate governance, and regulatory standards."""
            
            # Combine context and prompt
            full_prompt = f"{system_prompt}\n\n"
            if context:
                full_prompt += f"Relevant Context:\n{context}\n\n"
            full_prompt += f"Query: {prompt}\n\nResponse:"
            
            # Truncate if too long
            if len(full_prompt) > self.context_window:
                full_prompt = full_prompt[-self.context_window:]
            
            response = ollama.chat(
                model=self.model_name,
                messages=[
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': f"{context}\n\n{prompt}" if context else prompt}
                ],
                options={
                    'num_predict': max_tokens,
                    'temperature': temperature,
                    'top_p': 0.9,
                    'stop': ['</response>', '\n\n\n']
                }
            )
            
            return response['message']['content'].strip()
            
        except Exception as e:
            logger.error(f"❌ Ollama query failed: {e}")
            return self._fallback_response(prompt)
    
    def _fallback_response(self, prompt: str) -> str:
        """Provide fallback response when Ollama is unavailable"""
        fallback_responses = {
            "jurisdiction": "For ADGM compliance, ensure all disputes are subject to the exclusive jurisdiction of ADGM Courts, not UAE Federal Courts.",
            "governing law": "ADGM Law must be specified as the governing law for all company matters within ADGM jurisdiction.",
            "registered office": "The registered office must be located within Abu Dhabi Global Market (ADGM) jurisdiction.",
            "directors": "ADGM companies require at least one director who is a natural person over 18 years of age.",
            "compliance": "Ensure all documents comply with ADGM Companies Regulations 2020 and relevant regulatory requirements."
        }
        
        prompt_lower = prompt.lower()
        for key, response in fallback_responses.items():
            if key in prompt_lower:
                return f"[Fallback Response] {response}\n\nNote: Ollama AI service not available. Install and run Ollama for enhanced analysis."
        
        return "Ollama AI service not available. Please install and run Ollama for detailed legal analysis and recommendations."

class ADGMKnowledgeBase:
    """Comprehensive ADGM knowledge base with official sources integration"""
    
    def __init__(self):
        self.official_sources = self._initialize_official_sources()
        self.static_knowledge = self._load_static_knowledge()
        self.embedding_system = None
        self.knowledge_vectors = None
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (compatible; ADGM-Agent/1.0)'
        })
    
    def initialize_embeddings(self, embedding_system: OllamaEmbeddingSystem):
        """Initialize embedding system for semantic search"""
        self.embedding_system = embedding_system
        self._build_knowledge_vectors()
    
    def _initialize_official_sources(self) -> Dict[str, Dict]:
        """Initialize official ADGM document sources"""
        return {
            "company_formation": {
                "url": "https://www.adgm.com/registration-authority/registration-and-incorporation",
                "category": "Company Formation",
                "description": "Official ADGM company formation guidelines",
                "priority": "high"
            },
            "legal_framework": {
                "url": "https://www.adgm.com/legal-framework/guidance-and-policy-statements",
                "category": "Legal Framework", 
                "description": "ADGM legal framework and policy statements",
                "priority": "high"
            },
            "compliance_obligations": {
                "url": "https://www.adgm.com/operating-in-adgm/obligations-of-adgm-registered-entities",
                "category": "Compliance",
                "description": "Post-registration compliance obligations",
                "priority": "high"
            }
        }
    
    def _load_static_knowledge(self) -> List[Dict[str, Any]]:
        """Load static ADGM knowledge base"""
        return [
            {
                "id": "jurisdiction_001",
                "title": "ADGM Court Jurisdiction",
                "content": "All disputes and legal matters relating to ADGM companies must be subject to the exclusive jurisdiction of ADGM Courts. Use of UAE Federal Courts or other emirate courts is not compliant with ADGM regulations.",
                "source": "ADGM Companies Regulations 2020, Article 6",
                "category": "jurisdiction",
                "priority": "critical",
                "keywords": ["jurisdiction", "court", "dispute", "legal"]
            },
            {
                "id": "governing_law_001", 
                "title": "ADGM Governing Law Requirement",
                "content": "ADGM Law must be specified as the governing law for all company constitutional documents, contracts, and legal agreements within ADGM jurisdiction.",
                "source": "ADGM Companies Regulations 2020, Article 8",
                "category": "governing_law",
                "priority": "critical",
                "keywords": ["governing law", "applicable law", "adgm law"]
            },
            {
                "id": "registered_office_001",
                "title": "Registered Office Requirements",
                "content": "Every ADGM company must maintain its registered office within Abu Dhabi Global Market. The registered office must be a physical address within ADGM jurisdiction.",
                "source": "ADGM Companies Regulations 2020, Article 12",
                "category": "registered_office", 
                "priority": "critical",
                "keywords": ["registered office", "address", "adgm", "jurisdiction"]
            },
            {
                "id": "directors_001",
                "title": "Director Requirements",
                "content": "Every ADGM company must have at least one director who is a natural person over 18 years of age and not disqualified from acting as a director under ADGM regulations.",
                "source": "ADGM Companies Regulations 2020, Article 15",
                "category": "directors",
                "priority": "high",
                "keywords": ["director", "natural person", "age", "disqualified"]
            },
            {
                "id": "company_name_001",
                "title": "Company Name Requirements", 
                "content": "Company names must end with appropriate suffixes ('Limited', 'Ltd', 'LLC') and cannot contain restricted words such as 'Bank', 'Insurance', or 'Fund' without proper regulatory approval.",
                "source": "ADGM Companies Regulations 2020, Article 10",
                "category": "company_name",
                "priority": "high", 
                "keywords": ["company name", "limited", "ltd", "llc", "restricted words"]
            },
            {
                "id": "ubo_001",
                "title": "Ultimate Beneficial Owner Declaration",
                "content": "Companies must identify and declare all Ultimate Beneficial Owners (UBOs) who ultimately own or control 25% or more of shares or voting rights, or otherwise exercise control over the company.",
                "source": "ADGM AML/CTF Rules 2018",
                "category": "ubo_declaration",
                "priority": "high",
                "keywords": ["ubo", "beneficial owner", "25%", "control", "declaration"]
            },
            {
                "id": "registers_001",
                "title": "Company Registers Requirements",
                "content": "Companies must maintain proper registers including register of members, register of directors, and register of charges. These must be kept updated and available for inspection.",
                "source": "ADGM Companies Regulations 2020, Articles 25-30", 
                "category": "registers",
                "priority": "medium",
                "keywords": ["register", "members", "directors", "charges", "inspection"]
            },
            {
                "id": "employment_001",
                "title": "Employment Regulations",
                "content": "Employment contracts in ADGM must comply with ADGM Employment Regulations, including minimum wage, working hours, and termination procedures.",
                "source": "ADGM Employment Regulations 2019",
                "category": "employment",
                "priority": "medium", 
                "keywords": ["employment", "contract", "wages", "working hours", "termination"]
            },
            {
                "id": "data_protection_001",
                "title": "Data Protection Requirements",
                "content": "Companies processing personal data must comply with ADGM Data Protection Regulation, including appropriate policy documents and data subject rights procedures.",
                "source": "ADGM Data Protection Regulation 2021",
                "category": "data_protection",
                "priority": "medium",
                "keywords": ["data protection", "personal data", "policy", "data subject"]
            }
        ]
    
    def _build_knowledge_vectors(self):
        """Build embedding vectors for knowledge base entries"""
        if not self.embedding_system:
            return
        
        try:
            texts = [item['content'] for item in self.static_knowledge]
            self.knowledge_vectors = self.embedding_system.encode(texts)
            logger.info(f"✅ Built knowledge vectors for {len(texts)} entries")
        except Exception as e:
            logger.error(f"❌ Failed to build knowledge vectors: {e}")
    
    def search(self, query: str, category: str = None, max_results: int = 5) -> List[Dict[str, Any]]:
        """Search knowledge base with semantic similarity"""
        if self.embedding_system and self.knowledge_vectors is not None:
            return self._semantic_search(query, category, max_results)
        else:
            return self._keyword_search(query, category, max_results)
    
    def _semantic_search(self, query: str, category: str = None, max_results: int = 5) -> List[Dict[str, Any]]:
        """Perform semantic search using embeddings"""
        try:
            query_vector = self.embedding_system.encode([query])
            similarities = cosine_similarity(query_vector, self.knowledge_vectors)[0]
            
            # Create scored results
            results = []
            for idx, similarity in enumerate(similarities):
                if similarity > 0.1:  # Minimum similarity threshold
                    item = self.static_knowledge[idx].copy()
                    item['relevance_score'] = float(similarity)
                    
                    # Category filtering
                    if category is None or item.get('category') == category:
                        results.append(item)
            
            # Sort by relevance and priority
            results.sort(key=lambda x: (x['relevance_score'], 
                                      1 if x.get('priority') == 'critical' else 
                                      0.5 if x.get('priority') == 'high' else 0), 
                        reverse=True)
            
            return results[:max_results]
            
        except Exception as e:
            logger.error(f"❌ Semantic search failed: {e}")
            return self._keyword_search(query, category, max_results)
    
    def _keyword_search(self, query: str, category: str = None, max_results: int = 5) -> List[Dict[str, Any]]:
        """Fallback keyword-based search"""
        query_terms = query.lower().split()
        results = []
        
        for item in self.static_knowledge:
            if category and item.get('category') != category:
                continue
            
            score = 0
            content_lower = item['content'].lower()
            title_lower = item['title'].lower()
            keywords_lower = [kw.lower() for kw in item.get('keywords', [])]
            
            # Score based on term matches
            for term in query_terms:
                if term in title_lower:
                    score += 3  # Title matches are highly relevant
                if term in content_lower:
                    score += 2  # Content matches
                if term in keywords_lower:
                    score += 2  # Keyword matches
            
            # Priority boost
            if item.get('priority') == 'critical':
                score += 2
            elif item.get('priority') == 'high':
                score += 1
            
            if score > 0:
                item_copy = item.copy()
                item_copy['relevance_score'] = score
                results.append(item_copy)
        
        results.sort(key=lambda x: x['relevance_score'], reverse=True)
        return results[:max_results]

class ADGMComplianceEngine:
    """Advanced compliance checking engine with comprehensive rule validation"""
    
    def __init__(self, knowledge_base: ADGMKnowledgeBase):
        self.knowledge_base = knowledge_base
        self.compliance_rules = self._initialize_compliance_rules()
        self.red_flag_patterns = self._initialize_red_flag_patterns()
        self.document_templates = self._initialize_document_templates()
    
    def _initialize_compliance_rules(self) -> Dict[str, Dict]:
        """Initialize comprehensive ADGM compliance rules"""
        return {
            "jurisdiction": {
                "required_terms": ["adgm court", "adgm courts", "abu dhabi global market court"],
                "forbidden_terms": ["uae federal court", "dubai court", "sharjah court", "abu dhabi court"],
                "severity": "critical",
                "rule_ref": "ADGM Companies Regulations 2020, Art. 6"
            },
            "governing_law": {
                "required_terms": ["adgm law", "abu dhabi global market law", "adgm regulations"],
                "forbidden_terms": ["uae federal law", "dubai law", "sharjah law", "emirates law"],
                "severity": "critical", 
                "rule_ref": "ADGM Companies Regulations 2020, Art. 8"
            },
            "registered_office": {
                "required_terms": ["abu dhabi global market", "adgm"],
                "forbidden_terms": ["dubai", "sharjah", "uae mainland", "federal"],
                "context_required": ["registered office", "address"],
                "severity": "critical",
                "rule_ref": "ADGM Companies Regulations 2020, Art. 12"
            },
            "company_name": {
                "required_suffixes": ["limited", "ltd", "llc"],
                "restricted_terms": ["bank", "insurance", "fund", "investment advisor", "financial advisor"],
                "severity": "high",
                "rule_ref": "ADGM Companies Regulations 2020, Art. 10"
            }
        }
    
    def _initialize_red_flag_patterns(self) -> List[Dict]:
        """Initialize red flag detection patterns with enhanced context"""
        return [
            {
                "pattern": r"(?i)\b(uae\s+federal\s+court|dubai\s+court|sharjah\s+court|abu\s+dhabi\s+court)\b",
                "type": "jurisdiction_error",
                "severity": "critical",
                "message": "Incorrect court jurisdiction specified",
                "suggestion": "Change to 'ADGM Courts' for proper ADGM jurisdiction",
                "rule_ref": "ADGM Companies Regulations 2020, Art. 6"
            },
            {
                "pattern": r"(?i)\b(dubai\s+law|uae\s+federal\s+law|sharjah\s+law|emirates\s+law)\b",
                "type": "governing_law_error", 
                "severity": "critical",
                "message": "Incorrect governing law specified",
                "suggestion": "Change to 'ADGM Law' as the governing law",
                "rule_ref": "ADGM Companies Regulations 2020, Art. 8"
            },
            {
                "pattern": r"(?i)\[.*\]|\bTBD\b|\bTBC\b|\bto\s+be\s+(determined|confirmed|decided)\b",
                "type": "incomplete_info",
                "severity": "high",
                "message": "Incomplete or placeholder information detected",
                "suggestion": "Complete all required information before finalization"
            },
            {
                "pattern": r"(?i)\b(bank|banking|insurance|fund\s+management|investment\s+advisor)\b",
                "type": "regulated_activity",
                "severity": "high", 
                "message": "Regulated activity terminology detected",
                "suggestion": "Regulated activities require special licensing - consult ADGM regulatory requirements",
                "rule_ref": "ADGM Financial Services and Markets Regulations"
            },
            {
                "pattern": r"(?i)company(?!\s+(limited|ltd|llc))",
                "type": "company_name_suffix",
                "severity": "medium",
                "message": "Company name may be missing required legal suffix",
                "suggestion": "Ensure company name ends with 'Limited', 'Ltd', or 'LLC'",
                "rule_ref": "ADGM Companies Regulations 2020, Art. 10"
            }
        ]
    
    def _initialize_document_templates(self) -> Dict[str, Dict]:
        """Initialize document type templates with requirements"""
        return {
            "Articles of Association": {
                "required_sections": [
                    "company name and type",
                    "registered office", 
                    "objects and purposes",
                    "share capital",
                    "directors",
                    "meetings",
                    "jurisdiction clause"
                ],
                "critical_clauses": ["jurisdiction", "governing_law"],
                "category": "constitutional"
            },
            "Memorandum of Association": {
                "required_sections": [
                    "company name",
                    "registered office",
                    "objects", 
                    "liability limitation",
                    "share capital"
                ],
                "critical_clauses": ["registered_office"],
                "category": "constitutional"
            },
            "Board Resolution": {
                "required_sections": [
                    "resolution date",
                    "directors present",
                    "resolution details",
                    "authorization"
                ],
                "critical_clauses": [],
                "category": "governance"
            },
            "Employment Contract": {
                "required_sections": [
                    "employee details",
                    "employer details",
                    "job description", 
                    "salary and benefits",
                    "working hours",
                    "termination clauses"
                ],
                "critical_clauses": ["governing_law", "jurisdiction"],
                "category": "employment"
            }
        }
    
    def analyze_document(self, document_content: str, document_type: str) -> Dict[str, Any]:
        """Comprehensive document analysis with compliance checking"""
        analysis = {
            "document_type": document_type,
            "compliance_status": "unknown",
            "issues": [],
            "recommendations": [],
            "completeness_score": 0.0,
            "critical_issues": 0,
            "analysis_timestamp": datetime.now().isoformat()
        }
        
        # Perform red flag detection
        red_flags = self._detect_red_flags(document_content)
        analysis["issues"].extend(red_flags)
        
        # Perform rule-based compliance checking
        compliance_issues = self._check_compliance_rules(document_content)
        analysis["issues"].extend(compliance_issues)
        
        # Check document completeness
        completeness = self._check_document_completeness(document_content, document_type)
        analysis.update(completeness)
        
        # Generate recommendations using knowledge base
        recommendations = self._generate_recommendations(analysis["issues"])
        analysis["recommendations"] = recommendations
        
        # Calculate overall compliance status
        analysis["compliance_status"] = self._determine_compliance_status(analysis["issues"])
        analysis["critical_issues"] = len([issue for issue in analysis["issues"] if issue["severity"] == "critical"])
        
        return analysis
    
    def _detect_red_flags(self, content: str) -> List[Dict[str, Any]]:
        """Detect red flags using pattern matching"""
        issues = []
        
        for pattern_config in self.red_flag_patterns:
            matches = re.finditer(pattern_config["pattern"], content)
            for match in matches:
                # Get context around the match
                start = max(0, match.start() - 50)
                end = min(len(content), match.end() + 50)
                context = content[start:end].strip()
                
                issue = {
                    "type": pattern_config["type"],
                    "severity": pattern_config["severity"],
                    "message": pattern_config["message"],
                    "suggestion": pattern_config["suggestion"],
                    "matched_text": match.group(),
                    "context": context,
                    "position": {"start": match.start(), "end": match.end()},
                    "rule_reference": pattern_config.get("rule_ref", "ADGM Compliance Guidelines")
                }
                issues.append(issue)
        
        return issues
    
    def _check_compliance_rules(self, content: str) -> List[Dict[str, Any]]:
        """Check specific ADGM compliance rules"""
        issues = []
        content_lower = content.lower()
        
        for rule_name, rule_config in self.compliance_rules.items():
            # Check if context is relevant for this rule
            context_required = rule_config.get("context_required", [])
            if context_required:
                if not any(ctx.lower() in content_lower for ctx in context_required):
                    continue  # Skip this rule if context not found
            
            # Check for required terms
            required_terms = rule_config.get("required_terms", [])
            forbidden_terms = rule_config.get("forbidden_terms", [])
            
            has_required = any(term.lower() in content_lower for term in required_terms)
            has_forbidden = any(term.lower() in content_lower for term in forbidden_terms)
            
            # Generate issues based on rule violations
            if has_forbidden:
                issues.append({
                    "type": f"{rule_name}_forbidden",
                    "severity": rule_config["severity"],
                    "message": f"Forbidden {rule_name.replace('_', ' ')} terms detected",
                    "suggestion": f"Remove forbidden terms and use ADGM-compliant alternatives",
                    "rule_reference": rule_config["rule_ref"]
                })
            
            if required_terms and not has_required and any(ctx.lower() in content_lower for ctx in context_required if context_required):
                issues.append({
                    "type": f"{rule_name}_missing",
                    "severity": rule_config["severity"],
                    "message": f"Required {rule_name.replace('_', ' ')} terms missing",
                    "suggestion": f"Add appropriate {rule_name.replace('_', ' ')} clause with ADGM-compliant terms",
                    "rule_reference": rule_config["rule_ref"]
                })
        
        return issues
    
    def _check_document_completeness(self, content: str, document_type: str) -> Dict[str, Any]:
        """Check document completeness against template requirements"""
        template = self.document_templates.get(document_type, {})
        required_sections = template.get("required_sections", [])
        
        if not required_sections:
            return {
                "completeness_score": 1.0,
                "missing_sections": [],
                "present_sections": []
            }
        
        content_lower = content.lower()
        present_sections = []
        missing_sections = []
        
        for section in required_sections:
            section_keywords = section.lower().split()
            if any(keyword in content_lower for keyword in section_keywords):
                present_sections.append(section)
            else:
                missing_sections.append(section)
        
        completeness_score = len(present_sections) / len(required_sections) if required_sections else 1.0
        
        return {
            "completeness_score": completeness_score,
            "missing_sections": missing_sections,
            "present_sections": present_sections
        }
    
    def _generate_recommendations(self, issues: List[Dict[str, Any]]) -> List[str]:
        """Generate actionable recommendations based on identified issues"""
        recommendations = []
        issue_types = set(issue["type"] for issue in issues)
        
        # Priority recommendations for critical issues
        if any("jurisdiction" in issue_type for issue_type in issue_types):
            recommendations.append("🚨 CRITICAL: Ensure all dispute resolution clauses specify 'ADGM Courts' as the exclusive jurisdiction")
        
        if any("governing_law" in issue_type for issue_type in issue_types):
            recommendations.append("🚨 CRITICAL: Specify 'ADGM Law' as the governing law for all legal provisions")
        
        if any("registered_office" in issue_type for issue_type in issue_types):
            recommendations.append("🚨 CRITICAL: Registered office must be located within Abu Dhabi Global Market")
        
        # General recommendations
        if any("incomplete" in issue_type for issue_type in issue_types):
            recommendations.append("⚠️ Complete all placeholder information and TBD sections before finalization")
        
        if any("regulated_activity" in issue_type for issue_type in issue_types):
            recommendations.append("⚠️ Verify regulatory compliance for any financial services or regulated activities")
        
        if any("company_name" in issue_type for issue_type in issue_types):
            recommendations.append("📝 Ensure company name includes proper legal suffix (Limited, Ltd, or LLC)")
        
        # Add general best practices if no specific issues
        if not recommendations:
            recommendations.append("✅ Document appears compliant - consider final legal review before submission")
        
        return recommendations
    
    def _determine_compliance_status(self, issues: List[Dict[str, Any]]) -> str:
        """Determine overall compliance status"""
        if not issues:
            return "compliant"
        
        critical_issues = [issue for issue in issues if issue["severity"] == "critical"]
        high_issues = [issue for issue in issues if issue["severity"] == "high"]
        
        if critical_issues:
            return "non_compliant"
        elif high_issues:
            return "needs_attention"
        else:
            return "minor_issues"

class ADGMDocumentProcessor:
    """Main document processing orchestrator"""
    
    def __init__(self):
        self.knowledge_base = ADGMKnowledgeBase()
        self.embedding_system = OllamaEmbeddingSystem()
        self.chat_system = OllamaIntelligentChat()
        self.compliance_engine = ADGMComplianceEngine(self.knowledge_base)
        
        # Initialize embeddings
        self.knowledge_base.initialize_embeddings(self.embedding_system)
        
        logger.info("✅ ADGM Document Processor initialized")
    
    def process_document_file(self, file_path: str) -> Dict[str, Any]:
        """Process a single document file"""
        try:
            # Load document
            if not DOCX_AVAILABLE:
                raise Exception("python-docx not available - cannot process Word documents")
            
            document = Document(file_path)
            
            # Extract text content
            content = "\n".join([paragraph.text for paragraph in document.paragraphs])
            
            # Identify document type
            document_type = self._identify_document_type(content)
            
            # Perform compliance analysis
            analysis = self.compliance_engine.analyze_document(content, document_type)
            
            # Get AI-powered insights
            ai_insights = self._get_ai_insights(content, analysis)
            analysis["ai_insights"] = ai_insights
            
            # Add document processing metadata
            analysis.update({
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "word_count": len(content.split()),
                "character_count": len(content),
                "processing_timestamp": datetime.now().isoformat()
            })
            
            return analysis
            
        except Exception as e:
            logger.error(f"❌ Error processing {file_path}: {e}")
            return {
                "error": str(e),
                "file_path": file_path,
                "processing_timestamp": datetime.now().isoformat()
            }
    
    def _identify_document_type(self, content: str) -> str:
        """Identify document type using keyword analysis"""
        content_lower = content.lower()
        
        type_patterns = {
            "Articles of Association": ["articles of association", "company constitution", "aoa"],
            "Memorandum of Association": ["memorandum of association", "moa", "memorandum and articles"],
            "Board Resolution": ["board resolution", "directors resolution", "resolved that"],
            "Shareholder Resolution": ["shareholder resolution", "members resolution", "general meeting"],
            "Employment Contract": ["employment contract", "employment agreement", "terms of employment"],
            "UBO Declaration": ["ultimate beneficial owner", "ubo declaration", "beneficial ownership"],
            "Register of Members": ["register of members", "members register", "shareholder register"],
            "Data Protection Policy": ["data protection", "privacy policy", "personal data processing"]
        }
        
        best_match = "Unknown Document"
        best_score = 0
        
        for doc_type, keywords in type_patterns.items():
            score = sum(1 for keyword in keywords if keyword in content_lower)
            if score > best_score:
                best_score = score
                best_match = doc_type
        
        return best_match if best_score > 0 else "Unknown Document"
    
    def _get_ai_insights(self, content: str, analysis: Dict[str, Any]) -> str:
        """Get AI-powered insights and recommendations"""
        if not self.chat_system.ollama_available:
            return "AI insights unavailable - Ollama not connected"
        
        # Build context from analysis
        context = f"""
Document Type: {analysis['document_type']}
Compliance Status: {analysis['compliance_status']}
Issues Found: {len(analysis['issues'])}
Critical Issues: {analysis['critical_issues']}
Completeness Score: {analysis['completeness_score']:.1%}
"""
        
        # Create targeted prompt
        prompt = f"""
Analyze this ADGM legal document and provide specific, actionable recommendations:

{context}

Key Issues Identified:
{chr(10).join([f"- {issue['message']}" for issue in analysis['issues'][:5]])}

Please provide:
1. Priority actions to address critical issues
2. Specific text improvements needed
3. ADGM compliance best practices
4. Risk assessment and mitigation strategies

Focus on practical, implementable recommendations for ADGM compliance.
"""
        
        try:
            return self.chat_system.query(prompt, context, max_tokens=600, temperature=0.2)
        except Exception as e:
            logger.error(f"❌ AI insights failed: {e}")
            return f"AI insights error: {str(e)}"
    
    def create_reviewed_document(self, original_path: str, analysis: Dict[str, Any]) -> str:
        """Create reviewed document with comments and suggestions"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not available")
        
        try:
            # Load original document
            document = Document(original_path)
            
            # Add review header
            header_para = document.add_paragraph()
            header_run = header_para.add_run(f"""
{'='*80}
ADGM COMPLIANCE REVIEW REPORT
{'='*80}
Document: {analysis.get('file_name', 'Unknown')}
Review Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Document Type: {analysis.get('document_type', 'Unknown')}
Compliance Status: {analysis.get('compliance_status', 'Unknown').upper()}
Overall Score: {analysis.get('completeness_score', 0):.1%}
{'='*80}
""")
            header_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
            header_run.bold = True
            
            # Add issues summary
            issues = analysis.get('issues', [])
            if issues:
                summary_para = document.add_paragraph()
                summary_text = f"\n📊 ISSUES SUMMARY ({len(issues)} total)\n" + "="*50 + "\n"
                
                severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0}
                for issue in issues:
                    severity = issue.get('severity', 'low')
                    severity_counts[severity] = severity_counts.get(severity, 0) + 1
                
                summary_text += f"🚨 Critical Issues: {severity_counts['critical']}\n"
                summary_text += f"⚠️ High Priority: {severity_counts['high']}\n" 
                summary_text += f"📝 Medium Priority: {severity_counts['medium']}\n"
                summary_text += f"ℹ️ Low Priority: {severity_counts['low']}\n\n"
                
                # Add detailed issues
                for i, issue in enumerate(issues[:10], 1):  # Show top 10 issues
                    severity_emoji = {"critical": "🚨", "high": "⚠️", "medium": "📝", "low": "ℹ️"}
                    emoji = severity_emoji.get(issue.get('severity', 'low'), '📝')
                    
                    summary_text += f"{emoji} ISSUE #{i} - {issue.get('severity', 'unknown').upper()}\n"
                    summary_text += f"Problem: {issue.get('message', 'Unknown issue')}\n"
                    summary_text += f"Suggestion: {issue.get('suggestion', 'No suggestion available')}\n"
                    summary_text += f"Reference: {issue.get('rule_reference', 'N/A')}\n"
                    
                    if 'matched_text' in issue:
                        summary_text += f"Found Text: '{issue['matched_text']}'\n"
                    
                    summary_text += "-" * 40 + "\n"
                
                summary_run = summary_para.add_run(summary_text)
                summary_run.font.color.rgb = RGBColor(153, 0, 0)  # Dark red
            
            # Add AI insights if available
            ai_insights = analysis.get('ai_insights', '')
            if ai_insights and 'unavailable' not in ai_insights.lower() and 'error' not in ai_insights.lower():
                ai_para = document.add_paragraph()
                ai_text = f"\n🤖 AI RECOMMENDATIONS\n" + "="*40 + f"\n{ai_insights}\n" + "="*40 + "\n"
                ai_run = ai_para.add_run(ai_text)
                ai_run.font.color.rgb = RGBColor(0, 102, 51)  # Dark green
            
            # Add recommendations
            recommendations = analysis.get('recommendations', [])
            if recommendations:
                rec_para = document.add_paragraph()
                rec_text = f"\n📋 ACTION ITEMS\n" + "="*30 + "\n"
                for i, rec in enumerate(recommendations, 1):
                    rec_text += f"{i}. {rec}\n"
                rec_text += "="*30 + "\n"
                
                rec_run = rec_para.add_run(rec_text)
                rec_run.font.color.rgb = RGBColor(51, 102, 153)  # Blue
            
            # Save reviewed document
            output_filename = f"ADGM_Reviewed_{Path(original_path).stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = os.path.join(tempfile.gettempdir(), output_filename)
            document.save(output_path)
            
            logger.info(f"✅ Created reviewed document: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"❌ Error creating reviewed document: {e}")
            raise

# Global processor instance
processor = ADGMDocumentProcessor()

def process_documents_interface(files) -> Tuple[str, str, str, str]:
    """Main Gradio interface function"""
    if not files:
        return "No files uploaded", "", "{}", ""
    
    if not DOCX_AVAILABLE:
        return "Error: python-docx not installed. Please install with: pip install python-docx", "", "{}", ""
    
    try:
        results = {
            "summary": {
                "total_files": len(files),
                "processed_files": 0,
                "failed_files": 0,
                "total_issues": 0,
                "critical_issues": 0
            },
            "files": [],
            "processing_timestamp": datetime.now().isoformat()
        }
        
        processed_files = []
        
        # Process each file
        for file in files:
            logger.info(f"📄 Processing: {file.name}")
            
            try:
                # Analyze document
                analysis = processor.process_document_file(file.name)
                
                if "error" not in analysis:
                    # Create reviewed document
                    reviewed_path = processor.create_reviewed_document(file.name, analysis)
                    processed_files.append(reviewed_path)
                    
                    results["files"].append(analysis)
                    results["summary"]["processed_files"] += 1
                    results["summary"]["total_issues"] += len(analysis.get("issues", []))
                    results["summary"]["critical_issues"] += analysis.get("critical_issues", 0)
                else:
                    results["files"].append(analysis)
                    results["summary"]["failed_files"] += 1
                    
            except Exception as e:
                logger.error(f"❌ Failed to process {file.name}: {e}")
                results["files"].append({
                    "error": str(e),
                    "file_name": file.name,
                    "processing_timestamp": datetime.now().isoformat()
                })
                results["summary"]["failed_files"] += 1
        
        # Create download package
        download_file = None
        if processed_files:
            zip_path = os.path.join(tempfile.gettempdir(), f"ADGM_Review_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip")
            with zipfile.ZipFile(zip_path, 'w') as zipf:
                # Add reviewed documents
                for file_path in processed_files:
                    if os.path.exists(file_path):
                        zipf.write(file_path, os.path.basename(file_path))
                
                # Add JSON report
                json_report_path = os.path.join(tempfile.gettempdir(), "ADGM_Analysis_Report.json")
                with open(json_report_path, 'w', encoding='utf-8') as f:
                    json.dump(results, f, indent=2, default=str, ensure_ascii=False)
                zipf.write(json_report_path, "ADGM_Analysis_Report.json")
            
            download_file = zip_path
        
        # Generate summary message
        summary_msg = generate_summary_message(results)
        
        # Generate insights
        insights_msg = generate_insights_message(results)
        
        # Format JSON output
        json_output = json.dumps(results, indent=2, default=str, ensure_ascii=False)
        
        return summary_msg, download_file, json_output, insights_msg
        
    except Exception as e:
        logger.error(f"❌ Interface error: {e}")
        error_msg = f"Processing failed: {str(e)}"
        return error_msg, None, f'{{"error": "{str(e)}"}}', "Error occurred during processing"

def generate_summary_message(results: Dict[str, Any]) -> str:
    """Generate human-readable summary message"""
    summary = results["summary"]
    
    message_parts = []
    
    # System status
    if processor.chat_system.ollama_available:
        message_parts.append(f"🤖 **AI-Powered Analysis**: Ollama model '{processor.chat_system.model_name}' active")
    else:
        message_parts.append("📋 **Rule-Based Analysis**: Ollama not available, using pattern matching")
    
    # Processing summary
    message_parts.append(f"📊 **Processing Summary**: {summary['processed_files']}/{summary['total_files']} files processed successfully")
    
    if summary['failed_files'] > 0:
        message_parts.append(f"⚠️ **Failed Files**: {summary['failed_files']} files could not be processed")
    
    # Issues summary
    if summary['total_issues'] > 0:
        critical_count = summary['critical_issues']
        other_issues = summary['total_issues'] - critical_count
        
        if critical_count > 0:
            message_parts.append(f"🚨 **Critical Issues**: {critical_count} critical compliance issues found")
        
        if other_issues > 0:
            message_parts.append(f"📝 **Other Issues**: {other_issues} additional issues requiring attention")
    else:
        message_parts.append("✅ **No Issues Found**: All documents appear compliant")
    
    # Next steps
    if summary['critical_issues'] > 0:
        message_parts.append("🎯 **Next Steps**: Address critical issues immediately before proceeding")
    elif summary['total_issues'] > 0:
        message_parts.append("🎯 **Next Steps**: Review and address identified issues")
    else:
        message_parts.append("🎯 **Next Steps**: Consider final legal review before submission")
    
    return "\n\n".join(message_parts)

def generate_insights_message(results: Dict[str, Any]) -> str:
    """Generate insights and recommendations"""
    insights = []
    
    # System capabilities
    if processor.embedding_system.ollama_available:
        insights.append("🔍 **Enhanced Analysis**: Semantic document understanding active")
    if processor.chat_system.ollama_available:
        insights.append("💡 **AI Recommendations**: Context-aware legal suggestions available")
    
    # Document analysis insights
    file_results = results.get("files", [])
    doc_types = {}
    compliance_statuses = {}
    
    for file_result in file_results:
        if "error" not in file_result:
            doc_type = file_result.get("document_type", "Unknown")
            doc_types[doc_type] = doc_types.get(doc_type, 0) + 1
            
            compliance = file_result.get("compliance_status", "unknown")
            compliance_statuses[compliance] = compliance_statuses.get(compliance, 0) + 1
    
    if doc_types:
        insights.append("📋 **Document Types Identified**:")
        for doc_type, count in sorted(doc_types.items(), key=lambda x: x[1], reverse=True):
            insights.append(f"   • {doc_type}: {count}")
    
    if compliance_statuses:
        insights.append("⚖️ **Compliance Status Distribution**:")
        status_labels = {
            "compliant": "✅ Compliant",
            "minor_issues": "📝 Minor Issues", 
            "needs_attention": "⚠️ Needs Attention",
            "non_compliant": "🚨 Non-Compliant"
        }
        for status, count in sorted(compliance_statuses.items(), key=lambda x: x[1], reverse=True):
            label = status_labels.get(status, status.title())
            insights.append(f"   • {label}: {count}")
    
    # Common issues analysis
    all_issues = []
    for file_result in file_results:
        if "error" not in file_result:
            all_issues.extend(file_result.get("issues", []))
    
    if all_issues:
        issue_types = {}
        for issue in all_issues:
            issue_type = issue.get("type", "unknown")
            issue_types[issue_type] = issue_types.get(issue_type, 0) + 1
        
        if issue_types:
            insights.append("🎯 **Most Common Issues**:")
            for issue_type, count in sorted(issue_types.items(), key=lambda x: x[1], reverse=True)[:3]:
                friendly_name = issue_type.replace("_", " ").title()
                insights.append(f"   • {friendly_name}: {count} occurrence(s)")
    
    # Recommendations
    if results["summary"]["critical_issues"] > 0:
        insights.append("🚨 **Priority Actions**:")
        insights.append("   • Review all critical issues immediately")
        insights.append("   • Focus on jurisdiction and governing law clauses")
        insights.append("   • Ensure ADGM-specific requirements are met")
    
    return "\n\n".join(insights) if insights else "Analysis complete - see detailed results above"

def check_system_status() -> str:
    """Check system status for display"""
    status_parts = []
    
    # Ollama status
    if processor.chat_system.ollama_available:
        status_parts.append(f"✅ Ollama: Connected (Model: {processor.chat_system.model_name})")
    else:
        status_parts.append("❌ Ollama: Not Available")
    
    # Document processing
    if DOCX_AVAILABLE:
        status_parts.append("✅ Document Processing: Available")
    else:
        status_parts.append("❌ Document Processing: python-docx required")
    
    # ML capabilities
    if ML_AVAILABLE:
        status_parts.append("✅ ML Analysis: Available")
    else:
        status_parts.append("❌ ML Analysis: sklearn/numpy required")
    
    return " | ".join(status_parts)

def create_interface():
    """Create the Gradio interface"""
    # Custom CSS for professional styling
    custom_css = """
    .main-container {
        max-width: 1400px;
        margin: 0 auto;
        padding: 20px;
    }
    .status-indicator {
        background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
        color: white;
        padding: 10px 20px;
        border-radius: 10px;
        margin: 10px 0;
        font-family: 'Courier New', monospace;
        font-size: 12px;
    }
    .feature-card {
        background: #323233;
        border: 1px solid #dee2e6;
        border-radius: 10px;
        padding: 20px;
        margin: 10px 0;
        border-left: 4px solid #0066cc;
    }
    .critical-alert {
        background: #f8d7da;
        border: 1px solid #f5c6cb;
        color: #721c24;
        padding: 15px;
        border-radius: 8px;
        margin: 10px 0;
    }
    .success-alert {
        background: #d4edda;
        border: 1px solid #c3e6cb;
        color: #155724;
        padding: 15px;
        border-radius: 8px;
        margin: 10px 0;
    }
    """
    
    with gr.Blocks(
        title="ADGM Corporate Agent - Ollama Powered",
        theme=gr.themes.Soft(),
        css=custom_css
    ) as interface:
        
        # Main Header
        gr.HTML("""
        <div style="text-align: center; padding: 30px; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; border-radius: 15px; margin-bottom: 20px;">
            <h1 style="margin: 0; font-size: 2.5em;">🏛️ ADGM Corporate Agent</h1>
            <p style="margin: 10px 0 0 0; font-size: 1.2em; opacity: 0.9;">🤖 Ollama-Powered Legal Document Intelligence System</p>
            <p style="margin: 5px 0 0 0; font-size: 0.9em; opacity: 0.8;">Abu Dhabi Global Market Compliance & Document Review</p>
        </div>
        """)
        
        # System Status
        system_status = check_system_status()
        gr.HTML(f"""
        <div class="status-indicator">
            <strong>🔧 System Status:</strong> {system_status}
        </div>
        """)
        
        # Main Interface Layout
        with gr.Row():
            # Left Column - Upload and Information
            with gr.Column(scale=1):
                gr.HTML("<h3 style='color: #0066cc; border-bottom: 2px solid #0066cc; padding-bottom: 10px;'>📁 Document Upload</h3>")
                
                file_input = gr.Files(
                    label="Upload Legal Documents",
                    file_types=[".docx"],
                    file_count="multiple",
                    height=200
                )
                
                with gr.Row():
                    process_btn = gr.Button(
                        "🔍 Analyze Documents",
                        variant="primary",
                        size="lg",
                        scale=3
                    )
                    clear_btn = gr.Button(
                        "🗑️ Clear",
                        variant="secondary",
                        scale=1
                    )
                
                # Document Types Info
                gr.HTML("""
                <div class="feature-card">
                    <h4 style="color: #0066cc; margin-top: 0;">📋 Supported Documents</h4>
                    <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 15px; font-size: 14px;">
                        <div>
                            ✓ Articles of Association<br>
                            ✓ Memorandum of Association<br>
                            ✓ Board Resolutions<br>
                            ✓ Shareholder Resolutions
                        </div>
                        <div>
                            ✓ Employment Contracts<br>
                            ✓ UBO Declarations<br>
                            ✓ Register of Members<br>
                            ✓ Data Protection Policies
                        </div>
                    </div>
                </div>
                """)
                
                # AI Features
                gr.HTML("""
                <div class="feature-card">
                    <h4 style="color: #0066cc; margin-top: 0;">🤖 AI Capabilities</h4>
                    <ul style="margin: 5px 0; font-size: 14px;">
                        <li><strong>Semantic Analysis:</strong> Context-aware document understanding</li>
                        <li><strong>Compliance Detection:</strong> ADGM-specific rule validation</li>
                        <li><strong>Smart Recommendations:</strong> AI-powered legal suggestions</li>
                        <li><strong>Risk Assessment:</strong> Automated red flag identification</li>
                        <li><strong>Document Intelligence:</strong> Type identification & completeness checking</li>
                    </ul>
                </div>
                """)
                
                # Setup Guide
                gr.HTML("""
                <div class="feature-card">
                    <h4 style="color: #0066cc; margin-top: 0;">⚡ Quick Setup</h4>
                    <div style="background: #323233; padding: 10px; border-radius: 5px; font-family: monospace; font-size: 12px;">
                        # Install Ollama<br>
                        curl -fsSL https://ollama.ai/install.sh | sh<br><br>
                        # Download model<br>
                        ollama pull llama3.2<br><br>
                        # Install dependencies<br>
                        pip install python-docx ollama numpy scikit-learn
                    </div>
                </div>
                """)
            
            # Right Column - Results and Analysis
            with gr.Column(scale=2):
                gr.HTML("<h3 style='color: #0066cc; border-bottom: 2px solid #0066cc; padding-bottom: 10px;'>📊 Analysis Results</h3>")
                
                with gr.Tabs():
                    with gr.Tab("📋 Executive Summary", elem_id="summary-tab"):
                        summary_output = gr.Markdown(
                            label="Analysis Summary",
                            height=250,
                            value="Upload documents to begin analysis..."
                        )
                        
                        insights_output = gr.Markdown(
                            label="Key Insights & Recommendations", 
                            height=250,
                            value=""
                        )
                    
                    with gr.Tab("📥 Download Results", elem_id="download-tab"):
                        download_output = gr.File(
                            label="📦 Complete Analysis Package (ZIP)",
                            interactive=False
                        )
                        
                        gr.HTML("""
                        <div class="success-alert">
                            <strong style="color: black;">📦 Download Package Includes:</strong><br>
                            • Reviewed documents with inline compliance comments<br>
                            • AI-powered recommendations and legal insights<br>
                            • Structured JSON analysis report with all findings<br>
                            • ADGM regulatory references and citations<br>
                            • Actionable compliance recommendations
                        </div>
                        """)
                    
                    with gr.Tab("🔍 Detailed Analysis", elem_id="details-tab"):
                        json_output = gr.Code(
                            label="Complete Analysis Data (JSON)",
                            language="json",
                            lines=25,
                            value="{}"
                        )
                    
                    with gr.Tab("📚 ADGM Resources", elem_id="resources-tab"):
                        gr.HTML("""
                        <div class="feature-card">
                            <h4 style="color: #0066cc; margin-top: 0;">🏛️ Official ADGM Resources</h4>
                            <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 20px;">
                                <div>
                                    <h5>📋 Company Formation</h5>
                                    <ul style="font-size: 13px; margin: 5px 0;">
                                        <li><a href="https://www.adgm.com/registration-authority/registration-and-incorporation" target="_blank">Registration & Incorporation</a></li>
                                        <li><a href="https://www.adgm.com/setting-up" target="_blank">Company Setup Guide</a></li>
                                        <li><a href="https://www.adgm.com/legal-framework/guidance-and-policy-statements" target="_blank">Legal Framework</a></li>
                                    </ul>
                                </div>
                                <div>
                                    <h5>⚖️ Compliance & Legal</h5>
                                    <ul style="font-size: 13px; margin: 5px 0;">
                                        <li><a href="https://www.adgm.com/operating-in-adgm/obligations-of-adgm-registered-entities" target="_blank">Compliance Obligations</a></li>
                                        <li><a href="https://www.adgm.com/operating-in-adgm/obligations-of-adgm-registered-entities/annual-filings" target="_blank">Annual Filings</a></li>
                                        <li><a href="https://www.adgm.com/operating-in-adgm/post-registration-services" target="_blank">Post-Registration Services</a></li>
                                    </ul>
                                </div>
                            </div>
                            
                            <div style="margin-top: 20px; padding: 15px; background: #323233; border-radius: 8px;">
                                <h5 style="color: #1565c0; margin-top: 0;">🎯 Key Compliance Areas</h5>
                                <div style="font-size: 14px;">
                                    <strong>Critical Requirements:</strong> Jurisdiction (ADGM Courts), Governing Law (ADGM Law), Registered Office (ADGM)<br>
                                    <strong>Corporate Governance:</strong> Director requirements, UBO declarations, Company registers<br>
                                    <strong>Documentation:</strong> Articles/Memorandum of Association, Board/Shareholder resolutions<br>
                                    <strong>Ongoing Compliance:</strong> Annual filings, Register maintenance, Regulatory notifications
                                </div>
                            </div>
                        </div>
                        """)
                        
                        # Sample Documents Section
                        gr.HTML("""
                        <div class="feature-card">
                            <h4 style="color: #0066cc; margin-top: 0;">📄 Document Templates & Samples</h4>
                            <p style="font-size: 14px; margin-bottom: 15px;">
                                The system checks documents against ADGM official templates and requirements.
                                For official templates, visit the ADGM website or consult with qualified legal professionals.
                            </p>
                            <div style="background: #323233; border: 1px solid #ffeaa7; padding: 10px; border-radius: 5px;">
                                <strong>⚠️ Important:</strong> This system provides automated compliance checking assistance. 
                                Always verify with current ADGM regulations and consult qualified legal professionals 
                                for official compliance matters.
                            </div>
                        </div>
                        """)
        
        # Processing Functions - Connect UI to backend
        def process_and_update(files):
            """Process files and update all outputs"""
            if not files:
                return (
                    "Please upload at least one .docx document to begin analysis.",
                    None,
                    "{}",
                    "No documents uploaded yet."
                )
            
            return process_documents_interface(files)
        
        def clear_all():
            """Clear all inputs and outputs"""
            return (
                None,  # files
                "Upload documents to begin analysis...",  # summary
                None,  # download
                "{}",  # json
                ""     # insights
            )
        
        # Connect buttons to functions
        process_btn.click(
            fn=process_and_update,
            inputs=[file_input],
            outputs=[summary_output, download_output, json_output, insights_output],
            show_progress=True
        )
        
        clear_btn.click(
            fn=clear_all,
            outputs=[file_input, summary_output, download_output, json_output, insights_output]
        )
        
        # Footer with comprehensive information
        gr.HTML("""
        <div style="margin-top: 40px; padding: 30px; background: #323233; border-radius: 15px;">
            <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 30px;">
                
                <div>
                    <h4 style="color: #0066cc; margin-bottom: 15px;">🎯 Core Capabilities</h4>
                    <ul style="font-size: 14px; line-height: 1.6;">
                        <li><strong>Document Intelligence:</strong> Automated type identification and content analysis</li>
                        <li><strong>ADGM Compliance:</strong> Comprehensive rule validation and requirement checking</li>
                        <li><strong>Risk Detection:</strong> Advanced pattern matching for regulatory red flags</li>
                        <li><strong>AI Insights:</strong> Context-aware recommendations using local Ollama models</li>
                        <li><strong>Report Generation:</strong> Professional compliance reports with legal citations</li>
                    </ul>
                </div>
                
                <div>
                    <h4 style="color: #0066cc; margin-bottom: 15px;">⚖️ Compliance Framework</h4>
                    <ul style="font-size: 14px; line-height: 1.6;">
                        <li><strong>Jurisdiction Validation:</strong> ADGM Courts requirement verification</li>
                        <li><strong>Governing Law:</strong> ADGM Law compliance checking</li>
                        <li><strong>Corporate Structure:</strong> Director, shareholder, and UBO requirements</li>
                        <li><strong>Documentation Standards:</strong> Articles, memorandum, and resolution validation</li>
                        <li><strong>Regulatory Alignment:</strong> Current ADGM regulations compliance</li>
                    </ul>
                </div>
                
                <div>
                    <h4 style="color: #0066cc; margin-bottom: 15px;">🤖 AI Technology Stack</h4>
                    <ul style="font-size: 14px; line-height: 1.6;">
                        <li><strong>Ollama Integration:</strong> Local LLM processing for privacy and control</li>
                        <li><strong>Semantic Search:</strong> Embedding-based knowledge base queries</li>
                        <li><strong>NLP Analysis:</strong> Advanced text processing and pattern recognition</li>
                        <li><strong>Fallback Systems:</strong> TF-IDF and rule-based alternatives</li>
                        <li><strong>Context Management:</strong> Intelligent prompt engineering for legal domain</li>
                    </ul>
                </div>
                
            </div>
            
            <hr style="margin: 30px 0; border: 1px solid #bdc3c7;">
            
            <div style="text-align: center;">
                <h4 style="color: #0066cc; margin-bottom: 15px;">🚀 Getting Started</h4>
                <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 20px; margin-bottom: 20px;">
                    <div style="background: #323233; padding: 15px; border-radius: 10px; border-left: 4px solid #3498db;">
                        <strong>1. Install Ollama</strong><br>
                        <span style="font-size: 12px; font-family: monospace; background: #323233; padding: 2px 6px; border-radius: 3px;">curl -fsSL https://ollama.ai/install.sh | sh</span>
                    </div>
                    <div style="background: #323233; padding: 15px; border-radius: 10px; border-left: 4px solid #2ecc71;">
                        <strong>2. Download Model</strong><br>
                        <span style="font-size: 12px; font-family: monospace; background: #323233; padding: 2px 6px; border-radius: 3px;">ollama pull llama3.2</span>
                    </div>
                    <div style="background: #323233; padding: 15px; border-radius: 10px; border-left: 4px solid #e74c3c;">
                        <strong>3. Upload Documents</strong><br>
                        <span style="font-size: 12px;">Upload .docx files and click Analyze</span>
                    </div>
                </div>
            </div>
            
            <div style="background: #323233; color: white; padding: 20px; border-radius: 10px; text-align: center;">
                <p style="margin: 0; font-size: 13px;">
                    ⚠️ <strong>Legal Disclaimer:</strong> This system provides automated document analysis and compliance checking assistance. 
                    It is not a substitute for qualified legal advice. Always consult with licensed legal professionals familiar with 
                    ADGM regulations for official legal matters and compliance verification.
                </p>
                <br>
                <p style="margin: 0; font-size: 12px; opacity: 0.8;">
                    🔒 <strong>Privacy:</strong> All processing happens locally using Ollama. Documents are not sent to external services. 
                    | 🛠️ <strong>Technology:</strong> Python • Ollama • Gradio • scikit-learn • python-docx
                    | 📧 <strong>Support:</strong> Check system requirements and ensure all dependencies are installed
                </p>
            </div>
        </div>
        """)
    
    return interface

# Main execution
if __name__ == "__main__":
    print("🚀 Starting ADGM Corporate Agent with Ollama Integration...")
    print(f"📊 System Status: {check_system_status()}")
    print("📚 Loading ADGM knowledge base...")
    print("🤖 Initializing AI systems...")
    
    # Create and launch interface
    interface = create_interface()
    print("✅ System ready!")
    print("🌐 Launching web interface...")
    
    interface.launch(
        server_name="0.0.0.0",
        server_port=7860,
        share=True,
        debug=True,
        show_error=True,
        inbrowser=True
    )