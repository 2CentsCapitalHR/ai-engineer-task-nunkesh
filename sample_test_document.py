#!/usr/bin/env python3
"""
Sample Test Document Generator for ADGM Corporate Agent
Creates test documents with intentional compliance issues for testing
"""

from docx import Document
from docx.shared import Inches
import os

def create_test_articles_of_association():
    """Create a test Articles of Association with compliance issues"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Articles of Association', 0)
    title.alignment = 1  # Center alignment
    
    # Company details
    doc.add_heading('Test Company Limited', level=1)
    
    # Add content with intentional issues for testing
    content_sections = [
        {
            "title": "1. Company Name and Registration",
            "content": """
The name of the company is Test Company Limited. The company is incorporated under the laws of Dubai and shall be registered in Dubai, UAE.

The registered office of the company shall be located in Dubai International Financial Centre, Dubai, UAE.
            """
        },
        {
            "title": "2. Objects and Powers", 
            "content": """
The objects for which the company is established are to carry on business as a general trading company and to engage in any lawful business activities.

The company shall have full power to carry on any business which may seem to it to be capable of being conveniently carried on in connection with the above objects.
            """
        },
        {
            "title": "3. Share Capital",
            "content": """
The authorized share capital of the company is AED 100,000 divided into 100,000 ordinary shares of AED 1 each.

The company may issue shares to be determined and may increase or reduce the share capital as may be decided by the board of directors.
            """
        },
        {
            "title": "4. Directors",
            "content": """
The company shall have a minimum of one (1) director and a maximum of ten (10) directors.

The first directors of the company shall be appointed by the subscribers to the memorandum of association and shall hold office until the first annual general meeting.
            """
        },
        {
            "title": "5. Meetings and Resolutions",
            "content": """
The company shall hold an annual general meeting once in every calendar year, provided that not more than fifteen (15) months shall elapse between successive annual general meetings.

All general meetings other than annual general meetings shall be called extraordinary general meetings.
            """
        },
        {
            "title": "6. Dispute Resolution",
            "content": """
Any dispute arising out of or in connection with these articles shall be subject to the exclusive jurisdiction of the Dubai Courts and shall be governed by UAE Federal Law.

Without prejudice to the above, the parties may seek to resolve disputes through arbitration in accordance with the DIFC arbitration rules.
            """
        }
    ]
    
    # Add all sections
    for section in content_sections:
        doc.add_heading(section["title"], level=2)
        doc.add_paragraph(section["content"].strip())
        doc.add_paragraph("")  # Add spacing
    
    # Save the document
    filename = "Test_Articles_of_Association_with_Issues.docx"
    doc.save(filename)
    print(f"✅ Created test document: {filename}")
    return filename

def create_test_memorandum_of_association():
    """Create a test Memorandum of Association"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Memorandum of Association', 0)
    title.alignment = 1
    
    # Company details
    doc.add_heading('Test Company Limited', level=1)
    
    content_sections = [
        {
            "title": "1. Name of Company",
            "content": "The name of the company is Test Company Limited."
        },
        {
            "title": "2. Registered Office",
            "content": "The registered office of the company is situated in Sharjah, United Arab Emirates."
        },
        {
            "title": "3. Objects",
            "content": """
The objects for which the company is established are:
a) To carry on the business of general trading
b) To engage in consultancy services
c) To undertake any other lawful business activities as may be determined by the directors
            """
        },
        {
            "title": "4. Liability",
            "content": "The liability of the members is limited by shares."
        },
        {
            "title": "5. Share Capital", 
            "content": "The authorized share capital of the company is TBD, divided into ordinary shares of nominal value to be determined."
        }
    ]
    
    for section in content_sections:
        doc.add_heading(section["title"], level=2)
        doc.add_paragraph(section["content"].strip())
        doc.add_paragraph("")
    
    filename = "Test_Memorandum_of_Association_with_Issues.docx"
    doc.save(filename)
    print(f"✅ Created test document: {filename}")
    return filename

def create_test_board_resolution():
    """Create a test Board Resolution with issues"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Board Resolution', 0)
    title.alignment = 1
    
    doc.add_paragraph("Test Company Limited")
    doc.add_paragraph("Resolution No: BR-2024-001")
    doc.add_paragraph("Date: [To be determined]")
    
    doc.add_heading('Resolution for Company Incorporation', level=2)
    
    content = """
WHEREAS, the subscribers wish to incorporate a company under the laws of UAE Federal jurisdiction;

NOW THEREFORE, BE IT RESOLVED that:

1. The incorporation of Test Company Limited is hereby approved.

2. The registered office shall be established in Abu Dhabi mainland.

3. The company shall be governed by UAE Federal Law and any disputes shall be subject to UAE Federal Courts.

4. The authorized share capital shall be determined at a later date.

5. This resolution shall be binding upon all parties without prejudice to any future amendments.

IN WITNESS WHEREOF, the undersigned directors have executed this resolution.

Directors:
[Name to be confirmed]
[Signature required]

Date: ___________
    """
    
    doc.add_paragraph(content.strip())
    
    filename = "Test_Board_Resolution_with_Issues.docx"
    doc.save(filename)
    print(f"✅ Created test document: {filename}")
    return filename

def create_test_employment_contract():
    """Create a test Employment Contract"""
    
    doc = Document()
    
    title = doc.add_heading('Employment Contract', 0)
    title.alignment = 1
    
    doc.add_paragraph("Between: Test Company Limited")
    doc.add_paragraph("And: [Employee Name TBD]")
    
    content_sections = [
        {
            "title": "1. Employment Details",
            "content": """
Position: Senior Consultant
Start Date: To be confirmed
Employment Type: Full-time permanent
            """
        },
        {
            "title": "2. Workplace",
            "content": "The employee shall work at the company's registered office in Dubai, UAE."
        },
        {
            "title": "3. Governing Law",
            "content": "This contract shall be governed by Dubai Employment Law and any disputes shall be resolved in Dubai Courts."
        },
        {
            "title": "4. Compensation",
            "content": "Salary: AED [Amount to be determined] per month, subject to applicable deductions."
        }
    ]
    
    for section in content_sections:
        doc.add_heading(section["title"], level=2)
        doc.add_paragraph(section["content"].strip())
        doc.add_paragraph("")
    
    filename = "Test_Employment_Contract_with_Issues.docx"
    doc.save(filename)
    print(f"✅ Created test document: {filename}")
    return filename

def main():
    """Generate all test documents"""
    print("📄 Creating ADGM Test Documents...")
    print("=" * 50)
    
    try:
        # Create test documents
        test_files = [
            create_test_articles_of_association(),
            create_test_memorandum_of_association(), 
            create_test_board_resolution(),
            create_test_employment_contract()
        ]
        
        print("\n🎯 Test Documents Created Successfully!")
        print("=" * 50)
        print("These documents contain intentional compliance issues to test the ADGM Corporate Agent:")
        print()
        
        for i, filename in enumerate(test_files, 1):
            print(f"{i}. {filename}")
        
        print("\n🔍 Expected Issues to be Detected:")
        print("• Incorrect jurisdiction (Dubai Courts instead of ADGM Courts)")
        print("• Wrong governing law (UAE Federal Law instead of ADGM Law)")
        print("• Incorrect registered office location")
        print("• Incomplete information (TBD fields)")
        print("• Non-binding language")
        
        print("\n📋 Next Steps:")
        print("1. Run the ADGM Corporate Agent: python app.py")
        print("2. Upload these test documents to verify the system works")
        print("3. Check that compliance issues are properly detected")
        
        return True
        
    except Exception as e:
        print(f"❌ Error creating test documents: {e}")
        return False

if __name__ == "__main__":
    main()